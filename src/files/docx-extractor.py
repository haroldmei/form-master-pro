import docx
import json
import os
import hashlib
from collections import OrderedDict

def extract_docx_content(docx_path):
    """
    Extract content from a DOCX file with multilingual support.
    
    Args:
        docx_path (str): Path to the DOCX file
        
    Returns:
        dict: Content structure with paragraphs, tables, etc.
    """
    doc = docx.Document(docx_path)
    content = {
        "filename": os.path.basename(docx_path),
        "paragraphs": [],
        "tables": []
    }
    
    # Dictionary to track seen paragraph content (for deduplication)
    seen_paragraphs = set()
    
    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and text not in seen_paragraphs:
            seen_paragraphs.add(text)
            content["paragraphs"].append({
                "text": text,
                "style": para.style.name
            })
    
    # Extract tables
    for i, table in enumerate(doc.tables):
        table_data = []
        seen_cells = set()
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                # Only add non-empty cells that haven't been seen before within this table
                if cell_text and cell_text not in seen_cells:
                    seen_cells.add(cell_text)
                    row_data.append(cell_text)
            if row_data:
                table_data.append(row_data)
                
        if table_data:
            content["tables"].append({
                "id": i+1,
                "data": table_data
            })
    
    return content

def docx_to_json(docx_path, output_path=None):
    """
    Convert DOCX content to JSON and save to file.
    
    Args:
        docx_path (str): Path to the DOCX file
        output_path (str, optional): Path to save the JSON output. 
                                    If None, uses the docx filename with .json extension.
    
    Returns:
        str: Path to the output JSON file
    """
    if output_path is None:
        base_name = os.path.splitext(docx_path)[0]
        output_path = f"{base_name}.json"
    
    content = extract_docx_content(docx_path)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(content, f, ensure_ascii=False, indent=2)
    
    return output_path

def process_multiple_docx(directory_path, output_directory=None):
    """
    Process multiple DOCX files in a directory.
    
    Args:
        directory_path (str): Path to directory containing DOCX files
        output_directory (str, optional): Directory to save JSON files. 
                                        If None, uses the same directory.
    
    Returns:
        list: Paths to the output JSON files
    """
    if output_directory is None:
        output_directory = directory_path
    
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    
    results = []
    
    for filename in os.listdir(directory_path):
        if filename.lower().endswith('.docx'):
            docx_path = os.path.join(directory_path, filename)
            output_path = os.path.join(output_directory, f"{os.path.splitext(filename)[0]}.json")
            
            try:
                result_path = docx_to_json(docx_path, output_path)
                results.append(result_path)
                print(f"Successfully processed: {filename}")
            except Exception as e:
                print(f"Error processing {filename}: {str(e)}")
    
    return results

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='Extract content from DOCX files to JSON with multilingual support')
    parser.add_argument('input', help='Path to DOCX file or directory containing DOCX files')
    parser.add_argument('--output', help='Path to output JSON file or directory', default=None)
    parser.add_argument('--batch', action='store_true', help='Process all DOCX files in the input directory')
    
    args = parser.parse_args()
    
    if args.batch:
        if not os.path.isdir(args.input):
            print(f"Error: {args.input} is not a directory")
            exit(1)
        
        output_paths = process_multiple_docx(args.input, args.output)
        print(f"Processed {len(output_paths)} files. Output saved to: {args.output or args.input}")
    else:
        if not os.path.isfile(args.input) or not args.input.lower().endswith('.docx'):
            print(f"Error: {args.input} is not a valid DOCX file")
            exit(1)
        
        output_path = docx_to_json(args.input, args.output)
        print(f"Output saved to: {output_path}")