import docx
import json
import os
import re
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.shared import qn
from docx.oxml import parse_xml
import argparse
from collections import OrderedDict
from docx.oxml.ns import qn as oxml_qn

def extract_images(doc, docx_file):
    """Extract images from the document and save them."""
    image_dir = os.path.splitext(docx_file)[0] + "_images"
    if not os.path.exists(image_dir):
        os.makedirs(image_dir)
    
    image_dict = {}
    image_counter = 0
    
    # Access the main document part
    document_part = doc.part
    
    # Get all the image parts
    for rel in document_part.rels.values():
        if rel.reltype == RT.IMAGE:
            image_counter += 1
            image_name = f"image_{image_counter}{os.path.splitext(rel.target_part.partname)[1]}"
            image_path = os.path.join(image_dir, image_name)
            
            with open(image_path, 'wb') as f:
                f.write(rel.target_part.blob)
            
            image_dict[rel.rId] = {
                "path": image_path,
                "filename": image_name
            }
    
    return image_dict, image_dir

def remove_tracked_changes(doc):
    """Remove tracked changes and accept all revisions in the document."""
    try:
        # Get XML tree
        document_element = doc._element
        
        # Remove tracked changes (deletions)
        for element in document_element.xpath(".//w:del", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
        
        # Accept insertions (remove the mark but keep the content)
        for element in document_element.xpath(".//w:ins", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
            # Extract all child elements
            children = list(element)
            parent = element.getparent()
            if parent is not None:
                # Insert children before the ins element
                index = parent.index(element)
                for i, child in enumerate(children):
                    parent.insert(index + i, child)
                # Remove the ins element
                parent.remove(element)
        
        # Remove revision information
        for element in document_element.xpath(".//*[@w:rsidR or @w:rsidRPr or @w:rsidDel or @w:rsidP]", 
                                            namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
            for attr_name in [oxml_qn('w:rsidR'), oxml_qn('w:rsidRPr'), 
                             oxml_qn('w:rsidDel'), oxml_qn('w:rsidP')]:
                if attr_name in element.attrib:
                    del element.attrib[attr_name]
    except Exception as e:
        print(f"Warning: Could not fully process tracked changes: {e}")

def extract_tables(doc):
    """Extract tables from the document."""
    tables_data = []
    
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Handle cell content (might include paragraphs, runs with formatting)
                cell_content = extract_paragraph_content(cell.paragraphs)
                row_data.append(cell_content)
            table_data.append(row_data)
        tables_data.append(table_data)
    
    return tables_data

def extract_text_formatting(run):
    """Extract formatting information from a run."""
    formatting = {}
    
    # Safe attribute access with getattr and default values
    formatting["bold"] = run.bold if hasattr(run, 'bold') and run.bold is not None else None
    formatting["italic"] = run.italic if hasattr(run, 'italic') and run.italic is not None else None
    formatting["underline"] = run.underline if hasattr(run, 'underline') and run.underline is not None else None
    
    # Using the correct attribute for strikethrough
    if hasattr(run, 'font') and hasattr(run.font, 'strike'):
        formatting["strike"] = run.font.strike
    elif hasattr(run, 'font') and hasattr(run.font, 'strikethrough'):
        formatting["strike"] = run.font.strikethrough
    else:
        formatting["strike"] = None
    
    # Font properties with safe access
    if hasattr(run, 'font'):
        if hasattr(run.font, 'size') and run.font.size is not None:
            formatting["font_size"] = run.font.size.pt if hasattr(run.font.size, 'pt') else str(run.font.size)
        
        formatting["font_name"] = run.font.name if hasattr(run.font, 'name') else None
        
        if hasattr(run.font, 'color') and run.font.color is not None:
            formatting["color"] = run.font.color.rgb if hasattr(run.font.color, 'rgb') else None
        
        if hasattr(run.font, 'highlight_color') and run.font.highlight_color is not None:
            formatting["highlight"] = str(run.font.highlight_color)
    
    # Style
    formatting["style"] = run.style.name if hasattr(run, 'style') and run.style is not None else None
    
    # Filter out None values
    return {k: v for k, v in formatting.items() if v is not None}

def extract_run_content(run, image_dict):
    """Extract content from a run, which could be text or an image."""
    if hasattr(run, '_r') and len(run._r.xpath('.//w:drawing')) > 0:
        # This run contains an image
        for drawing in run._r.xpath('.//w:drawing'):
            blip = drawing.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
            if blip:
                rel_id = blip[0].get(qn('r:embed'))
                if rel_id in image_dict:
                    return {"type": "image", "image_data": image_dict[rel_id]}
        # If we couldn't find the image data
        return {"type": "image", "image_data": None}
    else:
        # This is a text run
        formatting = extract_text_formatting(run)
        if hasattr(run, 'text') and run.text:
            return {"type": "text", "text": run.text, "formatting": formatting}
        else:
            return None

def extract_paragraph_content(paragraphs):
    """Extract content from paragraphs, including runs with formatting."""
    result = []
    
    for para in paragraphs:
        para_data = {
            "text": para.text if hasattr(para, 'text') else "",
            "style": para.style.name if hasattr(para, 'style') and para.style is not None else None,
            "alignment": str(para.alignment) if hasattr(para, 'alignment') and para.alignment is not None else None,
            "runs": []
        }
        
        if hasattr(para, 'runs'):
            for run in para.runs:
                run_content = extract_run_content(run, {})  # Placeholder for image dict
                if run_content:
                    para_data["runs"].append(run_content)
        
        result.append(para_data)
    
    return result

def extract_docx_content(docx_file):
    """Extract all content from a DOCX file."""
    doc = docx.Document(docx_file)
    
    # Remove tracked changes to get the latest version only
    remove_tracked_changes(doc)
    
    # Extract images first
    image_dict, image_dir = extract_images(doc, docx_file)
    
    # Document properties
    doc_properties = {
        "core_properties": {}
    }
    
    # Safely access core properties
    if hasattr(doc, 'core_properties'):
        cp = doc.core_properties
        if hasattr(cp, 'title') and cp.title:
            doc_properties["core_properties"]["title"] = cp.title
        if hasattr(cp, 'author') and cp.author:
            doc_properties["core_properties"]["author"] = cp.author
        if hasattr(cp, 'created') and cp.created:
            doc_properties["core_properties"]["created"] = cp.created.isoformat()
        if hasattr(cp, 'modified') and cp.modified:
            doc_properties["core_properties"]["modified"] = cp.modified.isoformat()
        if hasattr(cp, 'revision') and cp.revision:
            doc_properties["core_properties"]["revision"] = cp.revision
    
    # Document content
    content = []
    
    # Process paragraphs
    for para in doc.paragraphs:
        # Skip empty paragraphs
        if not para.text.strip() and not any(run.text.strip() for run in para.runs):
            continue
            
        para_data = {
            "type": "paragraph",
            "text": para.text,
            "style": para.style.name if hasattr(para, 'style') and para.style is not None else None,
            "alignment": str(para.alignment) if hasattr(para, 'alignment') and para.alignment is not None else None,
            "runs": []
        }
        
        for run in para.runs:
            run_content = extract_run_content(run, image_dict)
            if run_content:
                para_data["runs"].append(run_content)
        
        content.append(para_data)
    
    # Process tables
    tables = extract_tables(doc)
    for table in tables:
        content.append({"type": "table", "data": table})
    
    # Consolidate all data
    result = {
        "properties": doc_properties,
        "content": content,
        "images_directory": image_dir if image_dict else None
    }
    
    return result

def save_to_json(data, output_file):
    """Save extracted data to a JSON file with UTF-8 encoding."""
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def main():
    parser = argparse.ArgumentParser(description='Extract latest version content from DOCX files to JSON.')
    parser.add_argument('input_file', help='Path to input DOCX file')
    parser.add_argument('-o', '--output', help='Path to output JSON file')
    parser.add_argument('--keep-history', action='store_true', help='Keep revision history (default: extract only latest version)')
    
    args = parser.parse_args()
    
    input_file = args.input_file
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' does not exist.")
        return
    
    output_file = args.output or os.path.splitext(input_file)[0] + '.json'
    
    print(f"Extracting latest version content from '{input_file}'...")
    content = extract_docx_content(input_file)
    
    print(f"Saving content to '{output_file}'...")
    save_to_json(content, output_file)
    
    print("Done!")

if __name__ == "__main__":
    main()